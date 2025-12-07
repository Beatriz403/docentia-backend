"""
Servicio de Exportación de Documentos
Convierte Markdown a Word (.docx) con formato profesional
"""
import re
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os


class ExportService:
    """Servicio para exportar documentos a diferentes formatos"""
    
    @staticmethod
    def markdown_to_docx(
        contenido_markdown: str,
        titulo: str = "Documento DocentIA",
        metadatos: Optional[Dict[str, Any]] = None
    ) -> Document:
        """
        Convierte contenido Markdown a documento Word con formato profesional
        
        Args:
            contenido_markdown: Contenido en formato Markdown
            titulo: Título del documento
            metadatos: Metadatos adicionales (autor, fecha, etc.)
        
        Returns:
            Objeto Document de python-docx
        """
        doc = Document()
        metadatos = metadatos or {}
        
        # Configurar propiedades del documento
        doc.core_properties.title = titulo
        doc.core_properties.author = metadatos.get('autor', 'DocentIA')
        
        # Añadir título principal
        titulo_p = doc.add_heading(titulo, level=0)
        titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Procesar contenido Markdown
        ExportService._procesar_markdown(doc, contenido_markdown)
        
        # Añadir pie de página
        ExportService._añadir_pie_pagina(doc)
        
        return doc
    
    @staticmethod
    def _procesar_markdown(doc: Document, contenido: str):
        """Procesa contenido Markdown y lo añade al documento"""
        
        lineas = contenido.split('\n')
        i = 0
        
        while i < len(lineas):
            linea = lineas[i].strip()
            
            # Saltar líneas vacías múltiples
            if not linea:
                i += 1
                continue
            
            # Encabezados
            if linea.startswith('# '):
                doc.add_heading(linea[2:], level=1)
            elif linea.startswith('## '):
                doc.add_heading(linea[3:], level=2)
            elif linea.startswith('### '):
                doc.add_heading(linea[4:], level=3)
            
            # Listas con viñetas
            elif linea.startswith('- ') or linea.startswith('* '):
                texto = linea[2:]
                p = doc.add_paragraph(texto, style='List Bullet')
            
            # Listas numeradas
            elif re.match(r'^\d+\.', linea):
                texto = re.sub(r'^\d+\.\s*', '', linea)
                p = doc.add_paragraph(texto, style='List Number')
            
            # Texto en negrita
            elif '**' in linea:
                p = doc.add_paragraph()
                partes = re.split(r'\*\*(.+?)\*\*', linea)
                for idx, parte in enumerate(partes):
                    if idx % 2 == 0:  # Texto normal
                        p.add_run(parte)
                    else:  # Texto en negrita
                        p.add_run(parte).bold = True
            
            # Párrafo normal
            else:
                doc.add_paragraph(linea)
            
            i += 1
    
    @staticmethod
    def _añadir_pie_pagina(doc: Document):
        """Añade pie de página al documento"""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "Generado con DocentIA - Recupera tu tiempo"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Estilo del pie de página
        footer_run = footer_para.runs[0]
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    @staticmethod
    def guardar_temporal(doc: Document, nombre_archivo: str) -> str:
        """
        Guarda el documento en un archivo temporal
        
        Args:
            doc: Documento de python-docx
            nombre_archivo: Nombre del archivo
        
        Returns:
            Ruta del archivo temporal
        """
        temp_dir = tempfile.gettempdir()
        ruta_completa = os.path.join(temp_dir, nombre_archivo)
        doc.save(ruta_completa)
        return ruta_completa


# Instancia global del servicio
export_service = ExportService()